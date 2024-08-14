import os
import time
import pdfkit
from jinja2 import Template
from docx import Document
from bs4 import BeautifulSoup
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = 'OutPut/'


def create_output_dir():
    if not os.path.exists(OUTPUT):
        os.mkdir(OUTPUT)


def create_folder():
    folder = time.strftime("%Y_%m_%d/")
    full_folder = OUTPUT + folder
    if not os.path.exists(full_folder):
        os.mkdir(full_folder)

    return full_folder


def create_name(type_):
    name = time.strftime("%H_%M_%S")
    return f'{name}.{type_}'


def render_template(template_file, context):
    with open(template_file, 'r') as file:
        template = Template(file.read())
    return template.render(context)


def save_as_html(cd_list, filename):
    # Create the HTML content with embedded styles
    with open('static/css/styles.css', 'r') as style_file:
        style = style_file.read()

    context = {
        'cd_list': cd_list,
        'style': style
    }

    rendered_html = render_template('templates/output_template.html', context)

    # Save the HTML content to the file
    with open(filename, 'w') as f:
        f.write(rendered_html)
    print(filename, 'is saved!')


def save_as_pdf(cd_list, filename):
    try:
        with open('static/css/styles.css', 'r') as style_file:
            style = style_file.read()

        # Create a simple HTML template for the PDF
        context = {
            'cd_list': cd_list,
            'style': style
        }

        rendered_html = render_template('templates/output_template.html', context)

        # Save the rendered HTML as a PDF
        pdfkit.from_string(rendered_html, filename)
        print(f"Saved PDF as {filename}")
    except OSError as e:
        if "No wkhtmltopdf executable found" in str(e):
            # Display a user-friendly error message
            error_message = """
            Error saving the table as PDF: No wkhtmltopdf executable found.

            To resolve this issue, please install wkhtmltopdf:
            1. Go to https://wkhtmltopdf.org/downloads.html
            2. Download and install the appropriate version for your operating system.
            3. Ensure that wkhtmltopdf is added to your system's PATH.

            After installing, you should be able to save the table as a PDF.
            """
            print(error_message)
            # Optionally, you can also display this message on the web page if needed
        raise e


def add_html_to_docx(html, doc):
    soup = BeautifulSoup(html, 'html.parser')
    for element in soup.find_all(['p', 'table', 'tr', 'td', 'th', 'sup', 'sub']):
        if element.name == 'table':
            # Get max number of columns in any row
            rows = element.find_all('tr')
            max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
            doc.add_heading(element.find_all('caption')[0].get_text())
            table = doc.add_table(rows=0, cols=max_cols)
            table.style = 'Table Grid'

            for row in rows:
                row_cells = table.add_row().cells
                cells = row.find_all(['td', 'th'])
                for idx, cell in enumerate(cells):
                    cell_paragraph = row_cells[idx].paragraphs[0]
                    for content in cell.contents:
                        if isinstance(content, str):
                            cell_paragraph.add_run(content)
                        elif content.name == 'sup':
                            run = cell_paragraph.add_run(content.get_text())
                            run.font.superscript = True
                        elif content.name == 'sub':
                            run = cell_paragraph.add_run(content.get_text())
                            run.font.subscript = True
        elif element.name == 'p':
            p = doc.add_paragraph(element.get_text())
            p.style.font.size = Pt(12)  # Adjust font size as needed
        # elif element.name == 'sup':
        #     run = doc.add_paragraph().add_run(element.get_text())
        #     run.font.superscript = True
        # elif element.name == 'sub':
        #     run = doc.add_paragraph().add_run(element.get_text())
        #     run.font.subscript = True


def save_as_word(cd_list, filename):
    with open('static/css/styles.css', 'r') as style_file:
        style = style_file.read()
    context = {
        'cd_list': cd_list,
        'style': style
    }
    html_content = render_template('templates/output_template.html', context)
    doc = Document()
    add_html_to_docx(html_content, doc)
    doc.save(filename)


def save_as_excel(cd_list, filename):
    # Implement Excel saving logic here
    pass


def save(cd_list, type_='html', name=None):
    create_output_dir()
    filename = create_folder()
    if name is None:
        name = create_name(type_)
        print(name)
    if type_.lower() == 'html':
        save_as_html(cd_list, filename + name)
    elif type_.lower() == 'pdf':
        save_as_pdf(cd_list, filename + name)
    elif type_.lower() == 'docx':
        save_as_word(cd_list, filename + name)
    elif type_.lower() == 'xlsx':
        save_as_excel(cd_list, filename + name)
    else:
        raise ValueError(f'unknown type {type_}')
